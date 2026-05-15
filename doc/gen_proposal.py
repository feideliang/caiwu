"""Generate a well-formatted business proposal Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Helper functions ──
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2B579A')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F2F6FC')

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('客户维度毛利率变动分析')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('客户 → 销售产品 → 订单 三级下钻及毛利率变动归因')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 新功能方案说明（业务版）')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()  # spacing


# ══════════════════════════════════════════════
# 1. 业务背景
# ══════════════════════════════════════════════
h = doc.add_heading('一、业务背景与需求', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(
    '当前系统支持按公司整体和事业部维度查看经营数据，但无法回答以下业务问题：'
)

problems = [
    ('「某某客户的毛利率为什么涨了/跌了？」', '没有客户维度的分析入口，只能按事业部看'),
    ('「毛利率变化是产品卖的结构变了，还是产品本身利润变了？」', '后台已计算了结构影响和毛利率影响的拆解数据，但前端没有展示'),
    ('「某个客户买了哪些产品？每个产品贡献了多少毛利？」', '系统存了销售产品数据，但没有按客户→产品的下钻链路'),
    ('「能看到对应的订单明细吗？」', '有订单级数据，但没有从客户→产品→订单的完整路径'),
]

for question, status in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'▸ {question}')
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(f'   现状：{status}')
    p2.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 2. 目标场景
# ══════════════════════════════════════════════
h = doc.add_heading('二、目标业务场景', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
run = p.add_run('典型用户操作路径（以阿里云为例）：')
run.font.size = Pt(12)
run.bold = True

steps = [
    ('Step 1', '客户总览', '打开「客户分析」页面，看到所有客户的收入、毛利率排名。瀑布图展示各客户对整体毛利率变化的贡献度。'),
    ('Step 2', '定位异常客户', '发现阿里云毛利率同比上涨 10 个百分点，在表格中点击阿里云进入下一层。'),
    ('Step 3', '查看产品明细', '看到阿里云采购的所有销售产品列表（收入、毛利率）。系统自动拆解：\n    • 结构影响：今年多卖了哪些高毛利产品？\n    • 毛利率影响：每个产品本身的毛利率变了多少？'),
    ('Step 4', '追溯订单', '点击某个产品（如 S5735-L48P4X），查看该产品在阿里云的所有订单明细。'),
]

for step, title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{step}：{title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    doc.add_paragraph(desc)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 3. 方案总览
# ══════════════════════════════════════════════
h = doc.add_heading('三、方案总览', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph('新增一个「客户分析」页面，按三级下钻设计：')

add_formatted_table(doc,
    ['层级', '页面内容', '关键指标', '可执行操作'],
    [
        ('L1\n客户总览', '所有客户收入/毛利率排名\n瀑布图（毛利率变动拆解）\n收入分布饼图', '营业收入\n毛利额\n毛利率\n客户数', '点击客户行 → 下钻到 L2'),
        ('L2\n产品明细', '该客户下的销售产品列表\n产品收入排名柱状图', '"客户维度转为\n销售产品维度"\n营业收入、毛利额、毛利率', '点击产品行 → 下钻到 L3\n点击返回 → 回 L1'),
        ('L3\n订单明细', '该产品在该客户的订单列表', '订单号\n收入/毛利/毛利率\n合同号', '查看订单详情\n点击返回 → 回 L2'),
    ],
    col_widths=[3, 5, 4.5, 4.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 4. 已具备的数据基础
# ══════════════════════════════════════════════
h = doc.add_heading('四、现有数据基础', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph('以下数据系统中已有，无需额外采集：')

add_formatted_table(doc,
    ['需要的数据', '系统中是否存在', '数据来源'],
    [
        ('客户名称', '✅ 已有', '财务数据中的 tags 字段（如"华为科技""阿里云"）'),
        ('销售产品名称', '✅ 已有', '财务数据中的 tags 字段（如"企业网络-EN3000"）'),
        ('订单号', '✅ 已有', '财务数据中的 order_id 字段'),
        ('客户级收入/成本/毛利率', '✅ 已有', '现有指标计算服务已支持按客户维度聚合'),
        ('毛利率变动拆解\n（结构影响 + 毛利率影响）', '✅ 后台已计算\n前端未展示', '指标计算服务已输出该数据，需在前端新增瀑布图展示'),
    ],
    col_widths=[4, 4, 8]
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('\n结论：')
run.bold = True
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run = p.add_run('所有业务数据已经就绪，只需新增前端页面和少量后端接口参数即可实现。')
run.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 5. 页面功能详解
# ══════════════════════════════════════════════
h = doc.add_heading('五、页面功能详解', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# L1
h2 = doc.add_heading('L1 — 客户总览', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

features = [
    ('筛选栏', '期间（月度/季度/年度）、对比基准（同比/环比）、客户选择'),
    ('KPI 卡片', '营业收入、毛利额、毛利率、客户总数'),
    ('毛利率变动拆解瀑布图', '核心亮点。展示每个客户对整体毛利率变化的结构影响和毛利率影响，直观看出哪个客户、什么原因导致了毛利率变化'),
    ('收入分布饼图', '各客户收入占比一目了然'),
    ('毛利率对比柱状图', '各客户毛利率横向对比'),
    ('客户明细表', '客户名、收入(万元)、收入贡献度、毛利率、结构影响、毛利率影响、总影响。点击客户行下钻'),
]
for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)

# L2
h2 = doc.add_heading('L2 — 产品明细', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

features2 = [
    ('钻取导航', '显示当前路径「客户名 > 销售产品」，有"返回客户列表"按钮'),
    ('KPI 卡片', '营业收入、毛利额、毛利率、产品数（该客户下）'),
    ('产品收入排名柱状图', '该客户采购金额最大的 Top 产品'),
    ('产品明细表', '销售产品名、收入(万元)、收入贡献度、毛利率。点击产品行下钻'),
]
for title, desc in features2:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)

# L3
h2 = doc.add_heading('L3 — 订单明细', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

p = doc.add_paragraph('显示该客户该产品的具体订单列表，包含订单号、收入、毛利、毛利率、合同号等信息。')
features3 = [
    ('钻取导航', '「客户名 > 产品名 > 订单」，可逐级返回'),
    ('订单列表', '订单号、期间、金额、毛利率、合同号等'),
]
for title, desc in features3:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 6. AI 助手
# ══════════════════════════════════════════════
h = doc.add_heading('六、AI 智能助手的角色', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph(
    '页面右侧的 AI 助手将接入客户维度的分析数据，支持自然语言归因问答。'
)
examples = [
    '「阿里云毛利率为什么涨了？」 → AI 结合拆解数据自动归因',
    '「哪个客户对毛利率拖累最大？」 → AI 读取数据后给出排名和影响值',
    '「企业网络-EN3000 在阿里云的订单有哪些异常？」 → AI 定位到具体订单',
]
for ex in examples:
    p = doc.add_paragraph()
    run = p.add_run(f'  例：{ex}')
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 7. 交付物
# ══════════════════════════════════════════════
h = doc.add_heading('七、交付物清单', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

add_formatted_table(doc,
    ['交付物', '说明'],
    [
        ('客户分析页面', '全新的「客户分析」页面，带三级下钻功能'),
        ('毛利率变动拆解瀑布图', '可视化展示结构影响和毛利率影响'),
        ('客户 → 产品 → 订单下钻', '点击式逐级下钻，操作直观'),
        ('AI 助手集成', 'AI 助手接入客户维度数据，支持归因问答'),
    ],
    col_widths=[5, 11]
)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 8. 价值总结
# ══════════════════════════════════════════════
h = doc.add_heading('八、业务价值总结', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

values = [
    ('归因能力', '从"毛利率变了"到"哪个客户、哪个产品导致的"，拉通归因链条'),
    ('可视化', '瀑布图直观展示结构和毛利率影响，业务人员一看就懂'),
    ('操作便捷', '点击式下钻，无需写 SQL 或提数'),
    ('数据驱动决策', '支持业务精准定位问题客户和产品，制定针对性策略'),
]

add_formatted_table(doc,
    ['价值点', '说明'],
    values,
    col_widths=[4, 12]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— END —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(11)

# ── Save ──
output_path = r'D:\workspace\caiwu04\doc\customer_analysis_proposal.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')