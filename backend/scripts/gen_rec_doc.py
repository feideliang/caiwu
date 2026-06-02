"""Generate business-friendly Word document for analysis recommendations."""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Load live data
with open('D:/workspace/caiwu04/rec_data.json', 'r', encoding='utf-8') as f:
    rec_data = json.load(f)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Style
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei UI'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei UI'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    return h

def add_para(text, bold=False, size=10, color=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Microsoft YaHei UI'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None, size=10):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = 'Microsoft YaHei UI'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei UI'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    return p

def format_val(val, key):
    if val is None:
        return '—'
    if 'margin' in key or 'concentration' in key or 'ratio' in key:
        return f'{val:.1f}%'
    if val > 100:
        return f'{val:,.0f}'
    return str(val)

# ════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('财务分析智能推荐系统')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
run.font.name = 'Microsoft YaHei UI'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 每个页面专属的智能分析指标与异常预警')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Microsoft YaHei UI'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('AI+BI 数智化财务管报系统 | 财务分析大模型能力说明')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Microsoft YaHei UI'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')

doc.add_page_break()

# ════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════
add_heading_styled('一、系统概述', level=1)

add_para('财务分析智能推荐系统是 AI+BI 数智化财务管报系统的核心能力之一。', size=11, space_after=8)

add_para('业务背景：', bold=True, size=11)
add_para('在传统的财务分析场景中，财务人员打开报表后需要自行判断“应该关注哪些指标”、“哪些数据异常”、“下一步应该往哪个方向深入分析”。这要求使用者具备丰富的财务分析经验，新手容易遗漏关键指标或忽略异常信号。', size=10.5)

add_para('系统能力：', bold=True, size=11)
add_para('本系统为每个财务分析页面自动生成专属的分析建议，包括：', size=10.5)

capabilities = [
    ('推荐关注指标', '系统根据当前页面类型和实际数据，自动推荐应重点关注的财务指标，并标注正常/预警/严重状态'),
    ('异常自动预警', '内置财务分析经验规则，当毛利率、收入同比等指标触及异常阈值时自动告警'),
    ('建议分析问题', '生成与当前页面数据高度相关的分析提问，一键即可启动深入分析'),
    ('建议下钻方向', '推荐下一步的分析维度，引导从宏观到微观的逐层深入分析'),
]
for title, desc in capabilities:
    add_bullet(desc, bold_prefix=f'{title}：')

doc.add_paragraph()
add_para('技术实现简述：', bold=True, size=11)
add_para('后端在 AI 分析服务层新增了一个“分析推荐引擎”。当用户打开某个分析页面时，系统会：', size=10.5)

for step in [
    '① 获取当前页面的核心财务数据（收入、毛利、客户集中度等）',
    '② 根据页面类型匹配对应的分析模板（如驾驶舱页面推荐宏观指标，产品页面推荐产品线毛利率）',
    '③ 将实际数据与内置的财务异常阈值进行比对，生成预警信息',
    '④ 将以上结果结构化返回前端，在智能助手面板中展示',
]:
    add_para(step, size=10.5)

doc.add_page_break()

# ════════════════════════════════════════
# 2. ANOMALY RULES
# ════════════════════════════════════════
add_heading_styled('二、异常检测规则', level=1)

add_para('以下规则来源于财务分析专家的实践经验，系统将其固化为自动检测能力：', size=10.5)

table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['异常指标', '预警阈值（黄色）', '严重阈值（红色）', '业务含义']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = 'Microsoft YaHei UI'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')

rows_data = [
    ['毛利率', '< 20%', '< 10%', '毛利率低于20%需常规审视，低于10%必须立即下钻分析'],
    ['毛利率', '> 60%', '—', '需确认是技术溢价还是偶然性项目，判断是否可复制'],
    ['单一客户占比', '> 10%', '—', '触发客户集中度风险预警，单一客户依赖度过高'],
    ['前三大客户集中度', '> 60%', '—', '客户集中度过高，需关注风险分散'],
    ['单一产品毛利占比', '> 40%', '—', '触发产品集中度风险预警'],
    ['收入同比', '< -10%', '—', '收入同比下滑超过10%，需深入分析原因'],
    ['负毛利订单占比', '> 10%', '> 15%', '亏损订单比例偏高，影响整体盈利能力'],
]

for r_idx, row_data in enumerate(rows_data, 1):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx].cells[c_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = 'Microsoft YaHei UI'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')

doc.add_paragraph()
add_para('注：以上阈值可在系统后台灵活调整，以适配不同行业、不同时期的业务特征。', size=9.5)

doc.add_page_break()

# ════════════════════════════════════════
# 3. PER-PAGE DETAILS
# ════════════════════════════════════════
add_heading_styled('三、各页面推荐建议（实时数据）', level=1)

add_para('以下为基于 2026年4月 实际数据，各页面自动生成的分析推荐：', size=10.5, space_after=10)

page_display = {
    '总览驾驶舱(Dashboard)': {
        'biz_desc': '公司整体经营情况的驾驶舱页面，面向管理层，展示收入、毛利、客户集中度等核心KPI',
    },
    '趋势分析': {
        'biz_desc': '展示收入、毛利、毛利率的时序变化趋势，用于判断增长/下滑的持续性',
    },
    '部门分析': {
        'biz_desc': '按销售部门（CBG/EBG/SBG等）拆解收入和毛利贡献，识别各业务线的盈利能力',
    },
    '产品分析': {
        'biz_desc': '按产品线分析收入、毛利和盈利能力，识别各产品线的风险与机会',
    },
    '客户分析': {
        'biz_desc': '按客户维度分析收入贡献和集中度，评估客户结构和依赖风险',
    },
    '核心指标(变动分析)': {
        'biz_desc': '收入/毛利额/毛利率的变动分析，四因素拆解（存续结构、存续毛利、新增、退出）',
    },
}

status_emoji = {'normal': '✅ 正常', 'warning': '⚠️ 预警', 'critical': '🔴 严重'}
severity_emoji = {'low': '🔵 提示', 'medium': '🟡 预警', 'high': '🔴 严重'}

for page_name, info in page_display.items():
    data = rec_data.get(page_name, {})
    if not data:
        continue

    add_heading_styled(page_name, level=2)

    add_para(info['biz_desc'], size=10, space_after=2)

    summary = data.get('summary', '')
    add_para(f'系统判断：{summary}', bold=True, size=10.5, color=(0x16, 0x77, 0xFF), space_after=6)

    if data.get('metrics'):
        add_para('推荐关注指标：', bold=True, size=10.5)
        mt = doc.add_table(rows=len(data['metrics'])+1, cols=4)
        mt.style = 'Light Grid Accent 1'
        mt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h_text in enumerate(['指标名称', '当前值', '状态', '分析建议']):
            cell = mt.rows[0].cells[i]
            cell.text = h_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = 'Microsoft YaHei UI'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
        for m_idx, m in enumerate(data['metrics'], 1):
            vals = [
                m['metric_name'],
                format_val(m.get('current_value'), m['metric_key']),
                status_emoji.get(m.get('status', 'normal'), m.get('status', '')),
                m.get('recommendation', ''),
            ]
            for c_idx, v in enumerate(vals):
                cell = mt.rows[m_idx].cells[c_idx]
                cell.text = v
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = 'Microsoft YaHei UI'
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
                if c_idx == 2:
                    if '预警' in v or '严重' in v:
                        set_cell_shading(cell, 'FFFBE6')
                    elif '正常' in v:
                        set_cell_shading(cell, 'F6FFED')

    if data.get('suggested_questions'):
        add_para('建议分析提问：', bold=True, size=10.5, space_after=2)
        for q in data['suggested_questions']:
            add_bullet(q, size=10)

    anom_count = len(data.get('anomalies', []))
    if anom_count > 0:
        add_para(f'异常告警（{anom_count}条）：', bold=True, size=10.5, color=(0xCF, 0x13, 0x22))
        for a in data['anomalies']:
            sev = severity_emoji.get(a.get('severity', 'low'), '')
            add_bullet(a['message'], bold_prefix=f'{sev} ', size=10)
    else:
        add_para('异常告警：当前无异常', size=10, color=(0x52, 0xC4, 0x1A))

    if data.get('drill_down_path'):
        add_para(f'建议下钻方向：{" → ".join(data["drill_down_path"])}', bold=True, size=10, color=(0x09, 0x58, 0xD9), space_after=2)

    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════
# 4. IMPLEMENTATION (business-friendly)
# ════════════════════════════════════════
add_heading_styled('四、实现原理（非技术版）', level=1)

add_para('本系统的核心思路是：将财务分析专家的经验规则与实时数据相结合，为每个分析页面自动生成“下一步该看什么”的建议。', size=10.5)

add_para('实现步骤：', bold=True, size=11)

for title, desc in [
    ('第一步：数据采集', '当用户打开某个分析页面时，系统会自动查询该页面的核心财务数据，包括营业收入、毛利额、毛利率、客户集中度、负毛利订单占比等指标。数据来源是系统中已汇总的财务聚合表。'),
    ('第二步：页面匹配', '系统根据页面类型（总览/趋势/部门/产品/客户/核心指标）匹配对应的分析模板。例如：总览页面推荐宏观三指标（收入、毛利率、客户集中度），产品页面推荐产品线毛利率和负毛利占比。'),
    ('第三步：异常检测', '将实际数据与内置的财务异常阈值逐条比对。比如当前毛利率为36.8%，高于20%预警线，标记为正常；如果某部门毛利率为-5%，低于10%严重阈值，则会生成红色告警。'),
    ('第四步：建议生成', '基于以上分析，系统生成四类输出：① 应关注的指标列表（带状态标识）；② 异常告警（如有）；③ 建议的分析提问；④ 建议的下一步下钻方向。'),
    ('第五步：前端展示', '在页面右侧的智能助手面板中，以结构化的方式展示以上推荐结果。用户可以直接点击建议问题，由 AI 财务助手给出详细分析回答。'),
]:
    add_bullet(desc, bold_prefix=f'{title}。', size=10)

doc.add_paragraph()

# ════════════════════════════════════════
# 5. AI ASSISTANT SYNERGY
# ════════════════════════════════════════
add_heading_styled('五、与 AI 财务助手的协同', level=1)

add_para('本系统不仅推荐“应该看什么”，还与 AI 财务助手无缝衔接，实现“看了之后怎么办”的闭环：', size=10.5)

for i, text in enumerate([
    '① 用户打开分析页面 → 系统自动推荐该页面的关注指标和分析问题',
    '② 用户点击推荐问题（如“毛利率为何偏低”）→ AI 财务助手基于实时数据和财务分析规则给出详细回答',
    '③ AI 回答中引用了具体的财务数据和业务规则（如毛利率公式、异常阈值）→ 回答有据可查',
    '④ 用户可继续追问深入细节 → AI 保持上下文，支持多轮对话分析',
]):
    add_para(text, size=10.5)

doc.add_paragraph()
add_para('示例交互流程：', bold=True, size=11)
add_para('用户在“产品分析”页面打开后，智能助手面板自动显示：', size=10.5)
add_bullet('产品线毛利率：—', size=10)
add_bullet('负毛利产品占比：—', size=10)
add_bullet('建议问题：为什么制造事业部毛利为负', size=10)
add_para('用户点击该问题后，AI 助手会基于当期数据，分析制造事业部的收入结构、成本构成、毛利率变化等因素，给出归因分析和建议。', size=10.5)

doc.add_paragraph()
add_para('—— 文档结束 ——', size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = 'D:/workspace/caiwu04/财务分析智能推荐系统说明文档.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
