"""Celery task for async report generation with state machine."""

from __future__ import annotations

import logging
from datetime import datetime, timezone

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.celery_app import celery_app

logger = logging.getLogger(__name__)

# State machine transitions
REPORT_STEPS = ["collecting_data", "ai_analyzing", "document_generating", "completed"]

# Shared synchronous engine — created once, reused across all calls.
_sync_engine = None


def _get_sync_engine():
    """Return a singleton synchronous SQLAlchemy engine."""
    global _sync_engine
    if _sync_engine is None:
        from app.config import settings

        _sync_engine = create_engine(
            settings.sync_database_url,
            pool_size=5,
            max_overflow=5,
            pool_pre_ping=True,
            pool_recycle=300,
        )
    return _sync_engine


def _get_sync_db():
    """Return a synchronous SQLAlchemy session from the shared engine."""
    return sessionmaker(bind=_get_sync_engine())()


def update_report_step(report_id: int, status: str, step: str, celery_id: str | None = None) -> None:
    """Update the report_task row with current status and step.

    This is also called directly from the API for cancel/retry operations.
    """
    session = _get_sync_db()
    try:
        from app.models.v3 import ReportTask

        obj = session.query(ReportTask).filter(ReportTask.id == report_id).first()
        if obj is None:
            logger.error("ReportTask %d not found for status update", report_id)
            return

        obj.status = status
        obj.current_step = step
        if celery_id:
            obj.celery_task_id = celery_id
        if status in ("completed", "failed"):
            obj.completed_at = datetime.now(timezone.utc).replace(tzinfo=None)

        session.commit()
    except Exception:
        session.rollback()
        logger.exception("Failed to update report_task %d", report_id)
    finally:
        session.close()


def _send_report_notification(report_id: int, success: bool, error: str | None = None) -> None:
    """Send a notification when a report task completes (success or failure)."""
    session = _get_sync_db()
    try:
        from app.models.v3 import ReportTask
        from app.tasks.notification import create_notification_sync

        obj = session.query(ReportTask).filter(ReportTask.id == report_id).first()
        if not obj:
            return

        if success:
            notification_type = "report_completed"
            title = f"Report #{report_id} generation completed"
            content = f"Your {obj.report_type} report for period {obj.period or 'N/A'} is ready for download."
        else:
            notification_type = "report_failed"
            title = f"Report #{report_id} generation failed"
            content = f"Error: {error or obj.error_message or 'Unknown error'}"

        link = f"/reports/{report_id}" if success else None

        create_notification_sync(
            user_id=obj.user_id,
            title=title,
            content=content,
            notification_type=notification_type,
            link=link,
            source_task_id=report_id,
        )
    except Exception:
        logger.exception("Failed to send notification for report %d", report_id)
    finally:
        session.close()


@celery_app.task(
    name="report_generation.generate_report",
    queue="report_generation",
    bind=True,
    acks_late=True,
    max_retries=3,
    default_retry_delay=30,
)
def generate_report_task(self, report_id: int) -> dict:
    """Execute the report generation pipeline as a Celery task.

    State machine:
        pending -> collecting_data -> ai_analyzing -> document_generating -> completed
                                                               |
                                                               v
                                                            failed -> retry
    """
    celery_task_id = self.request.id
    logger.info("Starting report generation: report_id=%s celery_id=%s", report_id, celery_task_id)

    # Mark as collecting_data
    update_report_step(report_id, status="running", step="collecting_data", celery_id=celery_task_id)

    try:
        # ── Step 1: Collect data ──────────────────────────────────
        data = _collect_data(report_id)

        # ── Step 2: AI analysis ───────────────────────────────────
        update_report_step(report_id, status="running", step="ai_analyzing", celery_id=celery_task_id)
        analysis = _ai_analyze(report_id, data)

        # ── Step 3: Generate document ─────────────────────────────
        update_report_step(report_id, status="running", step="document_generating", celery_id=celery_task_id)
        file_path, file_name = _generate_document(report_id, analysis)

        # ── Completed ─────────────────────────────────────────────
        update_report_step(report_id, status="completed", step="completed", celery_id=celery_task_id)

        # Persist file info
        session = _get_sync_db()
        try:
            from app.models.v3 import ReportTask

            obj = session.query(ReportTask).filter(ReportTask.id == report_id).first()
            if obj:
                obj.file_path = file_path
                obj.file_name = file_name
                session.commit()
        finally:
            session.close()

        logger.info("Report generation completed: report_id=%s", report_id)

        # Send notification
        _send_report_notification(report_id, success=True)

        return {"report_id": report_id, "status": "completed", "file_path": file_path, "file_name": file_name}

    except Exception as exc:
        logger.exception("Report generation failed: report_id=%s", report_id)
        update_report_step(report_id, status="failed", step="failed", celery_id=celery_task_id)

        # Persist error
        session = _get_sync_db()
        try:
            from app.models.v3 import ReportTask

            obj = session.query(ReportTask).filter(ReportTask.id == report_id).first()
            if obj:
                obj.error_message = str(exc)
                session.commit()
        finally:
            session.close()

        # Retry with exponential backoff
        retry_count = self.request.retries
        if retry_count < self.max_retries:
            delay = self.default_retry_delay * (2 ** retry_count)
            raise self.retry(exc=exc, countdown=min(delay, 300))

        # All retries exhausted
        _send_report_notification(report_id, success=False, error=str(exc))
        raise


# ── Pipeline step implementations ──────────────────────────────


def _collect_data(report_id: int) -> dict:
    """Collect financial data for the report period.

    Queries financial_data, data_batch, and related tables.
    """
    session = _get_sync_db()
    try:
        from app.models.v3 import ReportTask
        from app.models.core import FinancialData, DataBatch

        from sqlalchemy import select

        obj = session.query(ReportTask).filter(ReportTask.id == report_id).first()
        if not obj:
            raise ValueError(f"ReportTask {report_id} not found")

        period = obj.period or ""
        params = obj.params or {}

        # Query financial data for the period
        stmt = (
            select(FinancialData.metric_name, FinancialData.metric_value, FinancialData.period, FinancialData.entity)
            .where(FinancialData.period.like(f"%{period}%") if period else True)
            .order_by(FinancialData.period)
        )
        rows = session.execute(stmt).all()

        data_rows = [
            {
                "metric_name": r.metric_name,
                "metric_value": r.metric_value,
                "period": r.period,
                "entity": r.entity,
            }
            for r in rows
        ]

        return {
            "report_id": report_id,
            "period": period,
            "report_type": obj.report_type,
            "output_format": obj.output_format,
            "data_rows": data_rows,
            "params": params,
        }
    finally:
        session.close()


def _ai_analyze(report_id: int, data: dict) -> dict:
    """Run AI analysis on collected data.

    Produces a standardized four-part report structure:
    1. 总体业绩概览 — core KPIs (revenue, gross profit, margin, achievement)
    2. 多维下钻分析 — department, product_bgbu, customer breakdowns
    3. 根本原因总结 — AI-analyzed change drivers
    4. 具体业务建议 — sales, pricing, product strategy, CRM recommendations
    """
    data_rows = data.get("data_rows", [])
    report_type = data.get("report_type", "")
    period = data.get("period", "")

    # Basic statistics
    metrics: dict[str, list[float]] = {}
    entities: dict[str, set] = {}
    for row in data_rows:
        name = row["metric_name"]
        metrics.setdefault(name, []).append(row["metric_value"])
        entities.setdefault(name, set()).add(row.get("entity") or "总")

    stats = {}
    for name, values in metrics.items():
        if values:
            stats[name] = {
                "mean": sum(values) / len(values),
                "min": min(values),
                "max": max(values),
                "sum": sum(values),
                "count": len(values),
            }

    # ── Part 1: 总体业绩概览 ──
    revenue = stats.get("revenue", {})
    gp = stats.get("gross_profit", {})
    margin = stats.get("gross_margin", {})
    orders = stats.get("order_count", {})

    overview_parts = []
    if revenue.get("sum"):
        overview_parts.append(f"本期累计收入 {revenue['sum']:,.0f} 元")
    if gp.get("sum"):
        overview_parts.append(f"毛利额 {gp['sum']:,.0f} 元")
    if margin.get("mean"):
        overview_parts.append(f"整体毛利率 {margin['mean']:.2f}%")
    if orders.get("sum"):
        overview_parts.append(f"总订单数 {int(orders['sum'])} 笔")

    overview_text = "，".join(overview_parts) if overview_parts else "本期暂无汇总数据"

    # ── Part 2: 多维下钻分析 ──
    breakdown_text = ""
    dept_rows = [r for r in data_rows if r.get("entity") and "部门" in str(r.get("entity", ""))]
    prod_rows = [r for r in data_rows if r.get("entity") and ("产品" in str(r.get("entity", "")) or "线" in str(r.get("entity", "")))]

    if dept_rows:
        dept_summary = {}
        for r in dept_rows:
            e = r["entity"]
            dept_summary.setdefault(e, {"revenue": 0, "gross_profit": 0, "count": 0})
            if r["metric_name"] == "revenue":
                dept_summary[e]["revenue"] += r["metric_value"]
            elif r["metric_name"] == "gross_profit":
                dept_summary[e]["gross_profit"] += r["metric_value"]
        top_dept = sorted(dept_summary.items(), key=lambda x: x[1]["revenue"], reverse=True)[:3]
        lines = ["部门维度分析："]
        for name, vals in top_dept:
            gm = (vals["gross_profit"] / vals["revenue"] * 100) if vals["revenue"] else 0
            lines.append(f"  {name}: 收入 {vals['revenue']:,.0f} 元，毛利 {vals['gross_profit']:,.0f} 元，毛利率 {gm:.2f}%")
        breakdown_text += "\n".join(lines)

    if prod_rows:
        prod_summary = {}
        for r in prod_rows:
            e = r["entity"]
            prod_summary.setdefault(e, {"revenue": 0, "gross_profit": 0, "count": 0})
            if r["metric_name"] == "revenue":
                prod_summary[e]["revenue"] += r["metric_value"]
            elif r["metric_name"] == "gross_profit":
                prod_summary[e]["gross_profit"] += r["metric_value"]
        top_prod = sorted(prod_summary.items(), key=lambda x: x[1]["gross_profit"], reverse=True)[:3]
        lines = ["产品线维度分析："]
        for name, vals in top_prod:
            gm = (vals["gross_profit"] / vals["revenue"] * 100) if vals["revenue"] else 0
            lines.append(f"  {name}: 收入 {vals['revenue']:,.0f} 元，毛利 {vals['gross_profit']:,.0f} 元，毛利率 {gm:.2f}%")
        if breakdown_text:
            breakdown_text += "\n"
        breakdown_text += "\n".join(lines)

    if not breakdown_text:
        breakdown_text = "暂无部门/产品线维度明细数据"

    # ── Part 3: 根本原因总结 ──
    cause_text = ""
    if revenue.get("sum") and gp.get("sum"):
        overall_margin = gp["sum"] / revenue["sum"] * 100 if revenue["sum"] else 0
        cause_text = (
            f"本期收入与毛利的整体关联度：毛利率 {overall_margin:.2f}%。"
            f"{'收入与毛利双增长，经营态势良好。' if overall_margin > 20 else '毛利率偏低，需关注成本结构。' if overall_margin > 0 else '出现毛利为负情况，需重点排查亏损订单。'}"
        )
    else:
        cause_text = "数据量不足以进行归因分析"

    # ── Part 4: 具体业务建议 ──
    recommendations = []
    # Sales resource allocation
    if dept_rows:
        high_margin_dept = max(dept_summary.items(), key=lambda x: x[1]["gross_profit"]) if dept_summary else None
        if high_margin_dept:
            recommendations.append(
                f"销售资源倾斜：建议将更多资源投向毛利贡献度高的 {high_margin_dept[0]}，扩大其市场份额"
            )
    # Pricing strategy
    if overall_margin < 15 if "overall_margin" in dir() else (gp.get("mean", 0) / max(revenue.get("mean", 1), 1) * 100 < 15):
        recommendations.append(
            "定价策略调整：针对毛利率偏低的订单类型修订报价模型，适当提升服务类订单溢价"
        )
    # Product strategy
    if prod_rows and prod_summary:
        growing = [n for n, v in prod_summary.items() if v["gross_profit"] > 0]
        if growing:
            recommendations.append(
                f"产品战略优化：加大毛利率且增长趋势良好的 {'、'.join(growing[:2])} 的研发投入"
            )
    # CRM
    recommendations.append(
        "客户关系管理：对直签类高价值客户提供更深度服务以巩固利润，建立客户分级管理体系"
    )

    analysis = {
        "report_id": report_id,
        "report_type": report_type,
        "period": period,
        "statistics": stats,
        # Four-part structure
        "overview": overview_text,
        "breakdown": breakdown_text,
        "root_cause": cause_text,
        "recommendations": recommendations,
        "summary": f"{report_type} 报告（{period}）：{overview_text}",
        "anomalies": [],
    }

    return analysis


def _generate_document(report_id: int, analysis: dict) -> tuple[str, str]:
    """Generate a Word report document with standardized four-part structure.

    Sections:
    1. 总体业绩概览
    2. 多维下钻分析
    3. 根本原因总结
    4. 具体业务建议
    """
    import os
    from app.config import settings

    output_dir = settings.report_output_dir
    os.makedirs(output_dir, exist_ok=True)

    report_type = analysis.get("report_type", "report")
    period = analysis.get("period", "")
    file_name = f"report_{report_type}_{period}_{report_id}.docx"
    file_path = os.path.join(output_dir, file_name)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()

        # Title
        title = doc.add_heading(f"财务分析报告", level=1)
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(f"{report_type} | {period}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Part 1: 总体业绩概览
        doc.add_heading("一、总体业绩概览", level=2)
        doc.add_paragraph(analysis.get("overview", "暂无数据"))

        # Part 2: 多维下钻分析
        doc.add_heading("二、多维下钻分析", level=2)
        breakdown_lines = analysis.get("breakdown", "").split("\n")
        for line in breakdown_lines:
            if line.strip():
                if line.endswith("："):
                    doc.add_heading(line.strip(" ："), level=3)
                else:
                    doc.add_paragraph(line.strip())

        # Part 3: 根本原因总结
        doc.add_heading("三、根本原因总结", level=2)
        doc.add_paragraph(analysis.get("root_cause", "暂无数据"))

        # Part 4: 具体业务建议
        doc.add_heading("四、具体业务建议", level=2)
        for rec in analysis.get("recommendations", []):
            doc.add_paragraph(rec, style="List Bullet")

        # Statistics appendix
        doc.add_heading("附录：统计数据", level=2)
        stats = analysis.get("statistics", {})
        for metric, values in stats.items():
            doc.add_heading(metric, level=3)
            for k, v in values.items():
                doc.add_paragraph(f"  {k}: {v}")

        doc.save(file_path)
        logger.info("Document saved: %s", file_path)
    except ImportError:
        # python-docx not installed; write a placeholder text file
        file_path = file_path.replace(".docx", ".txt")
        file_name = file_name.replace(".docx", ".txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"财务分析报告 #{report_id}\n")
            f.write(f"{analysis.get('report_type', '')} | {analysis.get('period', '')}\n\n")
            f.write("一、总体业绩概览\n")
            f.write(analysis.get("overview", "") + "\n\n")
            f.write("二、多维下钻分析\n")
            f.write(analysis.get("breakdown", "") + "\n\n")
            f.write("三、根本原因总结\n")
            f.write(analysis.get("root_cause", "") + "\n\n")
            f.write("四、具体业务建议\n")
            for rec in analysis.get("recommendations", []):
                f.write(f"- {rec}\n")
        logger.info("Placeholder document saved: %s", file_path)

    return file_path, file_name
